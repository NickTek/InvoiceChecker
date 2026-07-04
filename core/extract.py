"""
core/extract.py

Handles turning uploaded files (PDF, Word, images) into plain text.
Tries the fast/direct method first, and falls back to OCR if that fails
or returns almost nothing (common with scanned documents).
"""

import os
import io


def extract_text(file_path: str) -> str:
    """
    Given a path to a file on disk, return its text content.
    Supports: .pdf, .docx, .doc, .txt, .png, .jpg, .jpeg
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    elif ext in (".png", ".jpg", ".jpeg"):
        return _extract_image(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_path: str) -> str:
    """Try direct text extraction first; fall back to OCR if it fails or is empty."""
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            pages_text = []
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                pages_text.append(page_text)
            text = "\n".join(pages_text).strip()
    except Exception:
        text = ""

    # If direct extraction failed or produced almost nothing, the PDF is
    # likely a scanned image -> fall back to OCR.
    if len(text) < 20:
        try:
            text = _ocr_pdf(file_path)
        except Exception as e:
            if not text:
                raise RuntimeError(
                    f"Could not read PDF using either direct extraction or OCR: {e}"
                )
    return text


def _ocr_pdf(file_path: str) -> str:
    """Convert PDF pages to images and run OCR on each page."""
    from pdf2image import convert_from_path
    import pytesseract

    images = convert_from_path(file_path)
    text_parts = []
    for image in images:
        text_parts.append(pytesseract.image_to_string(image))
    return "\n".join(text_parts).strip()


def _extract_docx(file_path: str) -> str:
    import docx
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    # Also pull text out of any tables (contracts often have pricing tables)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            paragraphs.append(row_text)
    return "\n".join(paragraphs).strip()


def _extract_image(file_path: str) -> str:
    import pytesseract
    from PIL import Image
    image = Image.open(file_path)
    return pytesseract.image_to_string(image).strip()
